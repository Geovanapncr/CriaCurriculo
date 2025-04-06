import streamlit as st
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from io import BytesIO


# Função para criar o currículo em .docx
def criar_curriculo(form):
    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add name as title (uppercase, bold, larger font)
    p = doc.add_paragraph()
    run = p.add_run(form['nome_completo'].upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Add contact information (simple, no bullets)
    contact_info = f"{form['telefone']} | {form['email']} | {form['endereco']} | {form['linkedin']}"
    p = doc.add_paragraph(contact_info)
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # Function to add horizontal line
    def add_horizontal_line():
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)  # Adjust the number of underscores as needed
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # Add OBJECTIVE section with lines
    add_horizontal_line()
    p = doc.add_paragraph("OBJETIVO")
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(14)
    p.style.font.bold = True
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(form['objetivo'])
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # Add EDUCATION section with lines
    add_horizontal_line()
    p = doc.add_paragraph("FORMAÇÕES")
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(14)
    p.style.font.bold = True
    p.paragraph_format.space_after = Pt(6)

    for formacao in form['formacoes']:
        # Bold course and institution
        p = doc.add_paragraph()
        run = p.add_run(
            f"{formacao['curso']} - {formacao['instituicao']} ({formacao['data_inicio']} - {formacao['data_fim']})")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Regular description with bullet
        p = doc.add_paragraph(formacao['descricao'], style='List Bullet')
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(12)
        p.paragraph_format.left_indent = Inches(0.25)

    # Add PROFESSIONAL EXPERIENCE if exists with lines
    if form.get('experiencias'):
        add_horizontal_line()
        p = doc.add_paragraph("EXPERIENCIA PROFISSIONAL")
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(14)
        p.style.font.bold = True
        p.paragraph_format.space_after = Pt(6)

        for experiencia in form['experiencias']:
            # Bold company and position
            p = doc.add_paragraph()
            run = p.add_run(
                f"{experiencia['empresa']} - {experiencia['cargo']} ({experiencia['data_inicio_emprego']} - {experiencia['data_fim_emprego']})")
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            # Regular description with bullet
            p = doc.add_paragraph(experiencia['descricao'], style='List Bullet')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.25)

    # Add QUALIFICATIONS if exists with lines
    if form.get('qualificacoes'):
        add_horizontal_line()
        p = doc.add_paragraph("QUALIFICAÇÕES")
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(14)
        p.style.font.bold = True
        p.paragraph_format.space_after = Pt(6)

        for qualificacao in form['qualificacoes']:
            p = doc.add_paragraph(qualificacao, style='List Bullet')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.25)

    # Add SKILLS if exists with lines
    if form.get('habilidades'):
        add_horizontal_line()
        p = doc.add_paragraph("HABILIDADES")
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(14)
        p.style.font.bold = True
        p.paragraph_format.space_after = Pt(6)

        for habilidade in form['habilidades']:
            p = doc.add_paragraph(habilidade, style='List Bullet')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.25)

    # Adiciona a nota de rodapé
    doc.add_paragraph("\n")  # Espaço antes do rodapé
    footer = doc.sections[0].footer
    p = footer.add_paragraph("Gerado por Geovana Panciera")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(10)
    p.style.font.color.rgb = RGBColor(0, 0, 0)  # Preto

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Título da aplicação
st.title("Gerador de Currículo")
st.write("Preencha seus dados abaixo e gere automaticamente um currículo profissional em Word.")

# Se não existir, inicialize as listas no session_state
if 'formacoes' not in st.session_state:
    st.session_state.formacoes = []

if 'experiencias' not in st.session_state:
    st.session_state.experiencias = []

if 'qualificacoes' not in st.session_state:
    st.session_state.qualificacoes = []

if 'habilidades' not in st.session_state:
    st.session_state.habilidades = []

# Formulário para preenchimento dos dados
with st.form("form_curriculo"):
    nome_completo = st.text_input("Nome completo", key="nome_completo")
    telefone = st.text_input("Telefone", key="telefone")
    email = st.text_input("E-mail", key="email")
    endereco = st.text_input("Endereço (Cidade, Estado)", key="endereco")
    linkedin = st.text_input("LinkedIn", key="linkedin")

    objetivo = st.text_area("Objetivo", "Ex: Atuar como jovem aprendiz.", key="objetivo")

    # Formatação da parte de formação
    st.subheader("Formação Acadêmica")
    curso = st.text_input("Curso", key="curso")
    instituicao = st.text_input("Instituição", key="instituicao")
    data_inicio = st.text_input("Data de Início", key="data_inicio")
    data_fim = st.text_input("Data de Finalização ou Previsão de Finalização",
                             key="data_fim")
    descricao = st.text_area("Descrição da Formação", key="descricao_formacao")

    # Botão de adicionar formação
    adicionar_formacao = st.form_submit_button("Adicionar Formação")

    if adicionar_formacao:
        # Adiciona a formação à lista se todos os campos estiverem preenchidos
        if curso and instituicao and data_inicio and descricao:
            st.session_state.formacoes.append({
                'curso': curso,
                'instituicao': instituicao,
                'data_inicio': data_inicio,
                'data_fim': data_fim or "Atualmente",  # Se não tiver data de fim, assume 'Atualmente'
                'descricao': descricao
            })
            # Limpa os campos para o usuário adicionar outra formação
            curso = ""
            instituicao = ""
            data_inicio = ""
            data_fim = ""
            descricao = ""
            st.success("Formação adicionada com sucesso!")
        else:
            st.error("Por favor, preencha todos os campos da formação.")

    # Botão de adicionar experiência profissional
    st.subheader("Experiência Profissional")
    empresa = st.text_input("Empresa", key="empresa")
    cargo = st.text_input("Cargo", key="cargo")
    data_inicio_emprego = st.text_input("Data de Início", key="data_inicio_emprego")
    data_fim_emprego = st.text_input("Data de Finalização (se for seu emprego atual coloque 'Atualmente')",
                                     key="data_fim_emprego")
    descricao_experiencia = st.text_area("Descrição do Trabalho", key="descricao_experiencia")

    adicionar_experiencia = st.form_submit_button("Adicionar Experiência")

    if adicionar_experiencia:
        # Adiciona a experiência à lista se todos os campos estiverem preenchidos
        if empresa and cargo and data_inicio_emprego and descricao_experiencia:
            st.session_state.experiencias.append({
                'empresa': empresa,
                'cargo': cargo,
                'data_inicio_emprego': data_inicio_emprego,
                'data_fim_emprego': data_fim_emprego or "Atualmente",  # Se não tiver data de fim, assume 'Atualmente'
                'descricao': descricao_experiencia
            })
            # Limpa os campos para o usuário adicionar outra experiência
            empresa = ""
            cargo = ""
            data_inicio_emprego = ""
            data_fim_emprego = ""
            descricao_experiencia = ""
            st.success("Experiência adicionada com sucesso!")
        else:
            st.error("Por favor, preencha todos os campos da experiência.")

    # Formulário para Qualificações e Habilidades
    st.subheader("Qualificações Profissionais")
    qualificacao = st.text_input("Qualificação", key="qualificacao")
    adicionar_qualificacao = st.form_submit_button("Adicionar Qualificação")

    if adicionar_qualificacao and qualificacao:
        st.session_state.qualificacoes.append(qualificacao)
        qualificacao = ""  # Limpar o campo após adicionar
        st.success("Qualificação adicionada com sucesso!")
    elif adicionar_qualificacao:
        st.error("Por favor, preencha o campo de qualificação.")

    st.subheader("Habilidades e Competências")
    habilidade = st.text_input("Habilidade", key="habilidade")
    adicionar_habilidade = st.form_submit_button("Adicionar Habilidade")

    if adicionar_habilidade and habilidade:
        st.session_state.habilidades.append(habilidade)
        habilidade = ""  # Limpar o campo após adicionar
        st.success("Habilidade adicionada com sucesso!")
    elif adicionar_habilidade:
        st.error("Por favor, preencha o campo de habilidade.")

    enviado = st.form_submit_button("Gerar Currículo")  # Botão de envio do formulário principal

    if enviado:
        dados = {
            'nome_completo': nome_completo,
            'telefone': telefone,
            'email': email,
            'endereco': endereco,
            'linkedin': linkedin,
            'objetivo': objetivo,
            'formacoes': st.session_state.formacoes,
            'experiencias': st.session_state.experiencias,
            'qualificacoes': st.session_state.qualificacoes,
            'habilidades': st.session_state.habilidades
        }
        arquivo = criar_curriculo(dados)

        st.success("Currículo gerado com sucesso!")

# Mover o botão de download para fora do formulário
if enviado:
    st.download_button(
        label="📄 Baixar Currículo em Word",
        data=arquivo,
        file_name="curriculo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
